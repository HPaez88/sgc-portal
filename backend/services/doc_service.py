"""
services/doc_service.py — Generación de documentos Word (.docx)
Basado en los formatos oficiales del SGC OOMAPASC de Cajeme.
"""
import io
import json
from datetime import datetime
from typing import Union

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────
# Helpers de estilo
# ─────────────────────────────────────────────
def _set_cell_bg(cell, hex_color: str):
    """Fondo de celda con color hex (sin #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _header_paragraph(doc: Document, text: str, size_pt: int = 11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p


def _section_title(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(f"  {text.upper()}  ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    # Dark blue background via table trick
    return p


def _field_row(doc: Document, label: str, value: str):
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.columns[0].width = Cm(6)
    t.columns[1].width = Cm(11)
    lc = t.cell(0, 0)
    lc.text = label
    lc.paragraphs[0].runs[0].bold = True
    lc.paragraphs[0].runs[0].font.size = Pt(9)
    _set_cell_bg(lc, "1F4E79")
    lc.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    vc = t.cell(0, 1)
    vc.text = value or "—"
    vc.paragraphs[0].runs[0].font.size = Pt(9)
    return t


def _add_dark_banner(doc: Document, text: str):
    """Agrega un párrafo con fondo azul oscuro simulado con tabla de 1 celda."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.cell(0, 0)
    _set_cell_bg(cell, "1F4E79")
    p = cell.paragraphs[0]
    run = p.add_run(f"  {text.upper()}  ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()


def _fmt_date(dt) -> str:
    if not dt:
        return "—"
    if isinstance(dt, datetime):
        return dt.strftime("%d/%m/%Y")
    return str(dt)


# ─────────────────────────────────────────────
# ACCIÓN CORRECTIVA
# ─────────────────────────────────────────────
def generar_accion_correctiva_docx(ac) -> bytes:
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # ── Encabezado ──
    t_header = doc.add_table(rows=2, cols=3)
    t_header.style = "Table Grid"
    # Fila 1
    t_header.cell(0, 0).merge(t_header.cell(1, 0)).text = "OOMAPASC"
    t_header.cell(0, 0).paragraphs[0].runs[0].bold = True
    t_header.cell(0, 0).paragraphs[0].runs[0].font.size = Pt(14)
    _set_cell_bg(t_header.cell(0, 0), "1F4E79")
    t_header.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    t_header.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    mid = t_header.cell(0, 1).merge(t_header.cell(1, 1))
    mid.text = "ACCIÓN CORRECTIVA\nSistema de Gestión de Calidad ISO 9001"
    mid.paragraphs[0].runs[0].bold = True
    mid.paragraphs[0].runs[0].font.size = Pt(12)
    mid.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_cell_bg(mid, "2E75B6")
    mid.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    t_header.cell(0, 2).text = f"Folio: {ac.folio}"
    t_header.cell(1, 2).text = f"Fecha apertura: {_fmt_date(ac.fecha_apertura)}"
    for r in [0, 1]:
        t_header.cell(r, 2).paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # ── DATOS GENERALES ──
    _add_dark_banner(doc, "Datos Generales")

    t_gen = doc.add_table(rows=2, cols=4)
    t_gen.style = "Table Grid"
    labels = ["Proceso:", "Área:", "Origen:", "Núm. Auditoría:"]
    values = [ac.proceso, ac.area, ac.origen, ac.num_auditoria or "N/A"]
    for i, (lbl, val) in enumerate(zip(labels, values)):
        row = i // 2
        col = (i % 2) * 2
        t_gen.cell(row, col).text = lbl
        t_gen.cell(row, col).paragraphs[0].runs[0].bold = True
        t_gen.cell(row, col).paragraphs[0].runs[0].font.size = Pt(9)
        _set_cell_bg(t_gen.cell(row, col), "D6E4F0")
        t_gen.cell(row, col + 1).text = val
        t_gen.cell(row, col + 1).paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # ── DESCRIPCIÓN DE LA NO CONFORMIDAD ──
    _add_dark_banner(doc, "Descripción de la No Conformidad")
    doc.add_paragraph(ac.descripcion_no_conformidad or "—").font.size = Pt(10)

    doc.add_paragraph(
        f"¿Impacta en otros procesos? {'SÍ' if ac.impacta_otros_procesos else 'NO'}"
    ).runs[0].bold = True
    if ac.impacta_otros_procesos and ac.procesos_afectados:
        doc.add_paragraph(f"Procesos afectados: {ac.procesos_afectados}")

    doc.add_paragraph()

    # ── EQUIPO DE TRABAJO ──
    _add_dark_banner(doc, "Equipo de Trabajo")
    try:
        equipo = json.loads(ac.equipo_trabajo or "[]")
    except Exception:
        equipo = []

    if equipo:
        t_eq = doc.add_table(rows=1 + len(equipo), cols=3)
        t_eq.style = "Table Grid"
        for j, hdr in enumerate(["Nombre", "Puesto", "Firma"]):
            t_eq.cell(0, j).text = hdr
            t_eq.cell(0, j).paragraphs[0].runs[0].bold = True
            _set_cell_bg(t_eq.cell(0, j), "2E75B6")
            t_eq.cell(0, j).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            t_eq.cell(0, j).paragraphs[0].runs[0].font.size = Pt(9)
        for i, miembro in enumerate(equipo, 1):
            t_eq.cell(i, 0).text = miembro.get("nombre", "")
            t_eq.cell(i, 1).text = miembro.get("puesto", "")
            t_eq.cell(i, 2).text = ""  # espacio para firma
            for j in range(3):
                t_eq.cell(i, j).paragraphs[0].runs[0].font.size = Pt(9)
    else:
        doc.add_paragraph("Sin miembros registrados.")

    doc.add_paragraph()

    # ── ANÁLISIS DE CAUSA RAÍZ ──
    _add_dark_banner(doc, "Análisis de Causa Raíz")

    if ac.accion_contenedora:
        p = doc.add_paragraph()
        p.add_run("Acción contenedora: ").bold = True
        p.add_run(ac.accion_contenedora).font.size = Pt(10)

    # Causas con puntuaciones
    try:
        causas = json.loads(ac.causas or "[]")
    except Exception:
        causas = []

    if causas:
        t_causas = doc.add_table(rows=1 + len(causas), cols=3)
        t_causas.style = "Table Grid"
        for j, hdr in enumerate(["#", "Causa", "Total / %"]):
            t_causas.cell(0, j).text = hdr
            t_causas.cell(0, j).paragraphs[0].runs[0].bold = True
            _set_cell_bg(t_causas.cell(0, j), "2E75B6")
            t_causas.cell(0, j).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            t_causas.cell(0, j).paragraphs[0].runs[0].font.size = Pt(9)
        for i, causa in enumerate(causas, 1):
            t_causas.cell(i, 0).text = str(i)
            t_causas.cell(i, 1).text = causa.get("causa", "")
            t_causas.cell(i, 2).text = str(causa.get("total", ""))
            for j in range(3):
                t_causas.cell(i, j).paragraphs[0].runs[0].font.size = Pt(9)

    if ac.causa_raiz_seleccionada:
        p = doc.add_paragraph()
        p.add_run("Causa raíz seleccionada: ").bold = True
        p.add_run(ac.causa_raiz_seleccionada)

    if ac.actualiza_matriz_riesgos:
        doc.add_paragraph(f"Actualiza matriz de riesgos: SÍ")
        if ac.descripcion_riesgo:
            doc.add_paragraph(f"Descripción: {ac.descripcion_riesgo}")

    doc.add_paragraph()

    # ── ACTIVIDADES ──
    _add_dark_banner(doc, "Actividades / Plan de Acción")
    try:
        actividades = json.loads(ac.actividades or "[]")
    except Exception:
        actividades = []

    if actividades:
        t_act = doc.add_table(rows=1 + len(actividades), cols=4)
        t_act.style = "Table Grid"
        for j, hdr in enumerate(["Actividad", "Responsable", "Fecha Término", "Evidencia"]):
            t_act.cell(0, j).text = hdr
            t_act.cell(0, j).paragraphs[0].runs[0].bold = True
            _set_cell_bg(t_act.cell(0, j), "2E75B6")
            t_act.cell(0, j).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            t_act.cell(0, j).paragraphs[0].runs[0].font.size = Pt(9)
        for i, act in enumerate(actividades, 1):
            t_act.cell(i, 0).text = act.get("actividad", "")
            t_act.cell(i, 1).text = act.get("responsable", "")
            t_act.cell(i, 2).text = act.get("fecha_termino", "")
            t_act.cell(i, 3).text = act.get("evidencia", "")
            for j in range(4):
                t_act.cell(i, j).paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # ── CIERRE ──
    _add_dark_banner(doc, "Cierre de la Acción")
    t_cierre = doc.add_table(rows=2, cols=2)
    t_cierre.style = "Table Grid"
    t_cierre.cell(0, 0).text = "Nombre y firma del auditor:"
    t_cierre.cell(0, 0).paragraphs[0].runs[0].bold = True
    t_cierre.cell(0, 1).text = ac.nombre_auditor_cierre or "________________________________"
    t_cierre.cell(1, 0).text = "Fecha de cierre:"
    t_cierre.cell(1, 0).paragraphs[0].runs[0].bold = True
    t_cierre.cell(1, 1).text = _fmt_date(ac.fecha_cierre)
    for r in range(2):
        for c in range(2):
            t_cierre.cell(r, c).paragraphs[0].runs[0].font.size = Pt(9)

    # Guardar en buffer
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────
# PLAN DE MEJORA
# ─────────────────────────────────────────────
def generar_plan_mejora_docx(pm) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # ── Encabezado ──
    t_header = doc.add_table(rows=2, cols=3)
    t_header.style = "Table Grid"

    logo_cell = t_header.cell(0, 0).merge(t_header.cell(1, 0))
    logo_cell.text = "OOMAPASC"
    logo_cell.paragraphs[0].runs[0].bold = True
    logo_cell.paragraphs[0].runs[0].font.size = Pt(14)
    _set_cell_bg(logo_cell, "1F4E79")
    logo_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    logo_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    mid = t_header.cell(0, 1).merge(t_header.cell(1, 1))
    mid.text = "PLAN DE MEJORA\nSistema de Gestión de Calidad ISO 9001"
    mid.paragraphs[0].runs[0].bold = True
    mid.paragraphs[0].runs[0].font.size = Pt(12)
    mid.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_cell_bg(mid, "2E75B6")
    mid.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    t_header.cell(0, 2).text = f"Folio: {pm.folio}"
    t_header.cell(1, 2).text = f"Fecha elaboración: {_fmt_date(pm.fecha_elaboracion)}"
    for r in [0, 1]:
        t_header.cell(r, 2).paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # ── DATOS GENERALES ──
    _add_dark_banner(doc, "Datos Generales")

    gen_data = [
        ("Título de la mejora:", pm.titulo_mejora),
        ("Gerencia / Coordinación:", pm.gerencia_coordinacion),
        ("Categoría de Mejora:", pm.categoria_mejora),
        ("Período de mejora:", pm.periodo_mejora),
        ("Origen:", pm.origen),
        ("Responsable:", pm.responsable),
        ("Integrantes:", pm.integrantes),
        ("Fecha cierre estimada:", _fmt_date(pm.fecha_cierre_estimada)),
    ]

    for lbl, val in gen_data:
        t = doc.add_table(rows=1, cols=2)
        t.style = "Table Grid"
        lc = t.cell(0, 0)
        lc.text = lbl
        lc.paragraphs[0].runs[0].bold = True
        lc.paragraphs[0].runs[0].font.size = Pt(9)
        _set_cell_bg(lc, "D6E4F0")
        vc = t.cell(0, 1)
        vc.text = str(val) if val else "—"
        vc.paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # ── DESCRIPCIÓN DE LA SITUACIÓN ACTUAL ──
    _add_dark_banner(doc, "Descripción de la Situación Actual")
    doc.add_paragraph(pm.descripcion_situacion_actual or "—")
    doc.add_paragraph()

    # ── SITUACIÓN DESEADA ──
    _add_dark_banner(doc, "Situación Deseada")
    doc.add_paragraph(pm.situacion_deseada or "—")
    doc.add_paragraph()

    # ── BENEFICIOS ──
    _add_dark_banner(doc, "Beneficios")
    doc.add_paragraph(pm.beneficios or "—")
    doc.add_paragraph()

    # ── ACTIVIDADES ──
    _add_dark_banner(doc, "Actividades del Plan")
    try:
        actividades = json.loads(pm.actividades or "[]")
    except Exception:
        actividades = []

    if actividades:
        headers = ["Actividad", "Responsable", "Indicador de Progreso",
                   "Fecha Término", "1er. Replanteo", "2do. Replanteo", "Evidencia"]
        t_act = doc.add_table(rows=1 + len(actividades), cols=len(headers))
        t_act.style = "Table Grid"
        for j, hdr in enumerate(headers):
            t_act.cell(0, j).text = hdr
            t_act.cell(0, j).paragraphs[0].runs[0].bold = True
            _set_cell_bg(t_act.cell(0, j), "2E75B6")
            t_act.cell(0, j).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            t_act.cell(0, j).paragraphs[0].runs[0].font.size = Pt(8)
        for i, act in enumerate(actividades, 1):
            row_vals = [
                act.get("actividad", ""),
                act.get("responsable", ""),
                act.get("indicador", ""),
                act.get("fecha_termino", ""),
                act.get("primer_replanteo", ""),
                act.get("segundo_replanteo", ""),
                act.get("evidencia", ""),
            ]
            for j, val in enumerate(row_vals):
                t_act.cell(i, j).text = val
                t_act.cell(i, j).paragraphs[0].runs[0].font.size = Pt(8)
    else:
        doc.add_paragraph("Sin actividades registradas.")

    doc.add_paragraph()

    # ── FIRMAS DE APROBACIÓN ──
    _add_dark_banner(doc, "Firmas de Aprobación")
    t_firmas = doc.add_table(rows=3, cols=3)
    t_firmas.style = "Table Grid"
    firmas = ["Director del Área", "Sistema de Gestión de Calidad", "Auditor Asignado a Cierre"]
    for j, firma in enumerate(firmas):
        t_firmas.cell(0, j).text = firma
        t_firmas.cell(0, j).paragraphs[0].runs[0].bold = True
        t_firmas.cell(0, j).paragraphs[0].runs[0].font.size = Pt(9)
        _set_cell_bg(t_firmas.cell(0, j), "D6E4F0")
        t_firmas.cell(1, j).text = ""  # espacio firma
        t_firmas.cell(1, j).height = Cm(1.5)
        t_firmas.cell(2, j).text = "Nombre y firma"
        t_firmas.cell(2, j).paragraphs[0].runs[0].font.size = Pt(8)

    doc.add_paragraph()

    # ── CIERRE ──
    _add_dark_banner(doc, "Cierre del Plan")
    t_cierre = doc.add_table(rows=3, cols=2)
    t_cierre.style = "Table Grid"
    cierre_data = [
        ("Fecha de cierre real:", _fmt_date(pm.fecha_cierre_real)),
        ("Auditor que cierra:", pm.nombre_auditor_cierre or "________________________________"),
        ("1er. Replanteo:", _fmt_date(pm.primer_replanteo)),
    ]
    for i, (lbl, val) in enumerate(cierre_data):
        t_cierre.cell(i, 0).text = lbl
        t_cierre.cell(i, 0).paragraphs[0].runs[0].bold = True
        t_cierre.cell(i, 0).paragraphs[0].runs[0].font.size = Pt(9)
        _set_cell_bg(t_cierre.cell(i, 0), "D6E4F0")
        t_cierre.cell(i, 1).text = val or "—"
        t_cierre.cell(i, 1).paragraphs[0].runs[0].font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
